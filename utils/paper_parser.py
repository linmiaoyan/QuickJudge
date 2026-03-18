# -*- coding: utf-8 -*-
"""
从 Word 试卷/答案文档中解析：选择题数量、主观题区块、答案、作文题干等。
支持英语等学科常见版式，可配合大模型增强解析。
"""
import re
import os
from typing import List, Dict, Any, Optional, Tuple

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# 常见答案行模式
RE_ANSWER_LINE_DOT = re.compile(r'(\d+)\s*[\.．]\s*([A-Da-d])')   # 1. A  2. B
RE_ANSWER_RANGE = re.compile(r'(\d+)\s*[-~－～]\s*(\d+)\s*[：:]\s*([A-Da-d\s]+)')  # 1-5: ABCDD
RE_ANSWER_SINGLE = re.compile(r'^(\d+)\s*[\.．、]\s*([A-Da-d])\s*$')
RE_ANSWER_BLOCK = re.compile(r'([A-Da-d])\s*([A-Da-d])\s*([A-Da-d])\s*([A-Da-d])\s*([A-Da-d])')  # 连续5个选项

# 题号/大题标题
RE_SECTION_TITLE = re.compile(
    r'^(第一部分?|第二部分?|第三部分?|第四部分?|第五部分?|'
    r'听力|阅读理解|阅读|完形填空|语法填空|短文改错|书面表达|写作|作文)'
    r'[：:\s]*.*?(\d+)\s*[-~－～]\s*(\d+)',
    re.I
)
RE_QUESTION_NUM = re.compile(r'^(\d+)\s*[\.．、]')  # 1. 或 1、


def _get_docx_text(doc: "DocumentType") -> Tuple[str, List[str], List[List[str]]]:
    """提取文档全文、段落列表、表格内容。"""
    paragraphs = []
    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if t:
            paragraphs.append(t)
    full_text = '\n'.join(paragraphs)
    tables_text = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [ (c.text or '').strip() for c in row.cells ]
            rows.append(cells)
        tables_text.append(rows)
    return full_text, paragraphs, tables_text


def _parse_choice_answers(paragraphs: List[str]) -> Dict[int, str]:
    """从段落中解析选择题答案，返回 {题号: 'A'|'B'|'C'|'D'}。"""
    answers = {}
    in_answer_section = False
    answer_keywords = ('答案', '参考答案', '选择题答案', '听力答案', '阅读答案')
    # 若全文大量出现 "数字. 选项" 则视为纯答案文档
    dot_matches = sum(len(RE_ANSWER_LINE_DOT.findall(p)) for p in paragraphs)
    if dot_matches >= 3 and not any(kw in '\n'.join(paragraphs[:5]) for kw in answer_keywords):
        in_answer_section = True

    for line in paragraphs:
        line = line.strip()
        if not line:
            continue
        # 进入答案区
        if any(kw in line and '答案' in line for kw in answer_keywords):
            in_answer_section = True
            # 同一行可能带答案，如 "答案：1. A 2. B"
            for m in RE_ANSWER_LINE_DOT.finditer(line):
                answers[int(m.group(1))] = m.group(2).upper()
            continue
        if not in_answer_section:
            continue
        # 1-5: ABCDD 形式
        for m in RE_ANSWER_RANGE.finditer(line):
            start, end = int(m.group(1)), int(m.group(2))
            opts = re.sub(r'\s+', '', m.group(3)).upper()
            for i, ch in enumerate(opts):
                if start + i <= end and ch in 'ABCD':
                    answers[start + i] = ch
        # 1. A  2. B 形式
        for m in RE_ANSWER_LINE_DOT.finditer(line):
            num, opt = int(m.group(1)), m.group(2).upper()
            if opt in 'ABCD':
                answers[num] = opt
        # 纯选项行：每 5 题一行等
        only_opts = re.sub(r'[\s\d\.．、]+', '', line)
        if only_opts and len(only_opts) <= 50 and all(c in 'ABCD' for c in only_opts.upper()):
            # 需要题号范围时从上下文推断，这里简单按顺序填（调用方可能已给 start）
            pass  # 暂不自动推断题号
    return answers


def _infer_choice_count_and_sections(paragraphs: List[str], full_text: str) -> List[Dict[str, Any]]:
    """
    推断选择题区块，如 听力 1-20、阅读 21-40。
    返回 [{"name": "听力", "start": 1, "end": 20}, ...]
    """
    sections = []
    # 常见英语试卷结构
    patterns = [
        (r'听力(?:\s*\(.*?\))?\s*[（(]?\s*(\d+)\s*[-~－～]\s*(\d+)', 1),
        (r'(?:第一部分?|Part\s*I)\s*[：:]\s*听力.*?(\d+)\s*[-~－～]\s*(\d+)', 1),
        (r'阅读理解(?:\s*\(.*?\))?\s*[（(]?\s*(\d+)\s*[-~－～]\s*(\d+)', 21),
        (r'阅读(?:\s*\(.*?\))?\s*[（(]?\s*(\d+)\s*[-~－～]\s*(\d+)', 21),
        (r'完形填空(?:\s*\(.*?\))?\s*[（(]?\s*(\d+)\s*[-~－～]\s*(\d+)', None),
        (r'七选五.*?(\d+)\s*[-~－～]\s*(\d+)', None),
    ]
    for pat, default_start in patterns:
        for m in re.finditer(pat, full_text, re.I):
            g = m.groups()
            if len(g) >= 2:
                try:
                    start = int(g[0])
                    end = int(g[1])
                    if start < end and end <= 100:
                        name = '选择题'
                        if '听力' in m.group(0) or 'Part' in m.group(0):
                            name = '听力'
                        elif '阅读' in m.group(0):
                            name = '阅读理解'
                        elif '完形' in m.group(0):
                            name = '完形填空'
                        sections.append({'name': name, 'start': start, 'end': end})
                except (ValueError, IndexError):
                    pass
    # 若未匹配到，根据答案推断总题数
    if not sections:
        # 从答案里找最大题号
        answers = _parse_choice_answers(paragraphs)
        if answers:
            max_num = max(answers.keys())
            sections.append({'name': '选择题', 'start': 1, 'end': max_num})
    return sections


def _infer_subjective_sections(paragraphs: List[str], full_text: str) -> List[Dict[str, Any]]:
    """
    推断主观题区块：语法填空、短文改错、书面表达等。
    返回 [{"name": "语法填空", "num_questions": 10, "num_lines": 15, "prompt": ""}, ...]
    """
    subjective = []
    # 语法填空：常为 10 空，约 15 行书写区
    if re.search(r'语法填空|（?:\d+）?\s*语法填空', full_text):
        subjective.append({
            'name': '语法填空',
            'num_questions': 10,
            'num_lines': 12,
            'prompt': ''
        })
    # 短文改错
    if re.search(r'短文改错|（?:\d+）?\s*短文改错', full_text):
        subjective.append({
            'name': '短文改错',
            'num_questions': 10,
            'num_lines': 12,
            'prompt': ''
        })
    # 书面表达/作文
    writing = re.search(
        r'(?:书面表达|写作|作文)(?:\s*[（(].*?[)）])?\s*[：:]\s*([^\n]+(?:\n(?!第).+)*)',
        full_text,
        re.S
    )
    if writing:
        prompt = writing.group(1).strip()[:1500]  # 题干/材料截断
        subjective.append({
            'name': '书面表达',
            'num_questions': 1,
            'num_lines': 20,
            'prompt': prompt
        })
    else:
        subjective.append({
            'name': '书面表达',
            'num_questions': 1,
            'num_lines': 20,
            'prompt': ''
        })
    return subjective


def parse_paper_docx(path: str, preset: Optional[str] = None) -> Dict[str, Any]:
    """
    解析 Word 试卷/答案文档。
    :param path: .docx 文件路径
    :param preset: 可选，学科/题型预设（如 'english'），后续可扩展其他预设以切换区块规则
    :return: {
        "choice_sections": [{"name","start","end"}],
        "choice_answers": {题号: "A"|"B"|"C"|"D"},
        "subjective_sections": [{"name","num_questions","num_lines","prompt"}],
        "full_text_preview": str (前 3000 字，供 LLM 用),
        "error": str (若有)
    }
    """
    if not DOCX_AVAILABLE:
        return {
            'choice_sections': [],
            'choice_answers': {},
            'subjective_sections': [],
            'full_text_preview': '',
            'error': '请安装 python-docx: pip install python-docx'
        }
    path = os.path.abspath(path)
    if not os.path.isfile(path) or not path.lower().endswith('.docx'):
        return {
            'choice_sections': [],
            'choice_answers': {},
            'subjective_sections': [],
            'full_text_preview': '',
            'error': '文件不存在或不是 .docx 格式'
        }
    try:
        doc = Document(path)
        full_text, paragraphs, tables_text = _get_docx_text(doc)
        choice_answers = _parse_choice_answers(paragraphs)
        choice_sections = _infer_choice_count_and_sections(paragraphs, full_text)
        if not choice_sections and choice_answers:
            max_num = max(choice_answers.keys()) if choice_answers else 0
            choice_sections = [{'name': '选择题', 'start': 1, 'end': max_num}]
        subjective_sections = _infer_subjective_sections(paragraphs, full_text)
        return {
            'choice_sections': choice_sections,
            'choice_answers': choice_answers,
            'subjective_sections': subjective_sections,
            'full_text_preview': full_text[:3000],
            'error': ''
        }
    except Exception as e:
        return {
            'choice_sections': [],
            'choice_answers': {},
            'subjective_sections': [],
            'full_text_preview': '',
            'error': str(e)
        }


def parse_paper_docx_with_llm(path: str, llm_callback: Optional[callable] = None, preset: Optional[str] = None) -> Dict[str, Any]:
    """
    先做基础解析，再用 LLM 从 full_text_preview 中抽取/修正：选择题区、主观题区、作文题干。
    llm_callback(full_text_preview) 应返回 JSON 或 dict，可包含：
    choice_sections, choice_answers, subjective_sections, essay_prompt 等覆盖/补充字段。
    preset 同 parse_paper_docx，用于后续按学科切换解析规则。
    """
    result = parse_paper_docx(path, preset=preset)
    if result.get('error'):
        return result
    if not llm_callback:
        return result
    try:
        extra = llm_callback(result['full_text_preview'])
        if isinstance(extra, str):
            import json
            extra = json.loads(extra)
        if isinstance(extra, dict):
            if 'choice_sections' in extra:
                result['choice_sections'] = extra['choice_sections']
            if 'choice_answers' in extra:
                result['choice_answers'].update(extra['choice_answers'])
            if 'subjective_sections' in extra:
                result['subjective_sections'] = extra['subjective_sections']
            if 'essay_prompt' in extra and result['subjective_sections']:
                for s in result['subjective_sections']:
                    if s.get('name') in ('书面表达', '写作', '作文'):
                        s['prompt'] = extra['essay_prompt']
                        break
    except Exception:
        pass
    return result
